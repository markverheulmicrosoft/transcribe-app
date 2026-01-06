"""
Export service for generating Word and PDF documents from transcriptions.
"""
import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

from app.services.speech_service import TranscriptionResult


class ExportService:
    """Service for exporting transcriptions to Word and PDF formats."""
    
    @staticmethod
    def create_word_document(result: TranscriptionResult) -> io.BytesIO:
        """
        Create a Word document from a transcription result.
        
        Args:
            result: The transcription result to export
            
        Returns:
            BytesIO buffer containing the Word document
        """
        doc = Document()
        
        # Add title
        title = doc.add_heading("Transcriptie", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add metadata
        doc.add_paragraph()
        meta_para = doc.add_paragraph()
        meta_para.add_run("Bestand: ").bold = True
        meta_para.add_run(result.filename)
        
        meta_para = doc.add_paragraph()
        meta_para.add_run("Datum: ").bold = True
        meta_para.add_run(result.created_at.strftime("%d-%m-%Y %H:%M"))
        
        meta_para = doc.add_paragraph()
        meta_para.add_run("Taal: ").bold = True
        meta_para.add_run(result.locale)
        
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        doc.add_paragraph()
        
        # Add transcription content
        doc.add_heading("Inhoud", level=1)
        
        current_speaker = None
        current_text = []
        
        for segment in result.segments:
            if segment.speaker_id != current_speaker:
                # Write previous speaker's content
                if current_text and current_speaker:
                    para = doc.add_paragraph()
                    speaker_run = para.add_run(f"{current_speaker}: ")
                    speaker_run.bold = True
                    para.add_run(" ".join(current_text))
                    doc.add_paragraph()
                
                current_speaker = segment.speaker_id
                current_text = [segment.text]
            else:
                current_text.append(segment.text)
        
        # Write last speaker's content
        if current_text and current_speaker:
            para = doc.add_paragraph()
            speaker_run = para.add_run(f"{current_speaker}: ")
            speaker_run.bold = True
            para.add_run(" ".join(current_text))
        
        # Add footer
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        footer = doc.add_paragraph()
        footer.add_run("Gegenereerd door Transcriptie PoC - Raad van State").italic = True
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    @staticmethod
    def create_pdf_document(result: TranscriptionResult) -> io.BytesIO:
        """
        Create a PDF document from a transcription result.
        
        Args:
            result: The transcription result to export
            
        Returns:
            BytesIO buffer containing the PDF document
        """
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm
        )
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1  # Center
        )
        
        speaker_style = ParagraphStyle(
            'Speaker',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=12,
            spaceBefore=12
        )
        
        meta_style = ParagraphStyle(
            'Meta',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#666666'
        )
        
        # Build content
        story = []
        
        # Title
        story.append(Paragraph("Transcriptie", title_style))
        story.append(Spacer(1, 20))
        
        # Metadata
        story.append(Paragraph(f"<b>Bestand:</b> {result.filename}", meta_style))
        story.append(Paragraph(
            f"<b>Datum:</b> {result.created_at.strftime('%d-%m-%Y %H:%M')}", 
            meta_style
        ))
        story.append(Paragraph(f"<b>Taal:</b> {result.locale}", meta_style))
        story.append(Spacer(1, 30))
        
        # Separator
        story.append(Paragraph("_" * 70, styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Content header
        story.append(Paragraph("Inhoud", styles['Heading2']))
        story.append(Spacer(1, 15))
        
        # Transcription content
        current_speaker = None
        current_text = []
        
        for segment in result.segments:
            if segment.speaker_id != current_speaker:
                if current_text and current_speaker:
                    text = " ".join(current_text)
                    story.append(Paragraph(
                        f"<b>{current_speaker}:</b> {text}",
                        speaker_style
                    ))
                
                current_speaker = segment.speaker_id
                current_text = [segment.text]
            else:
                current_text.append(segment.text)
        
        # Last speaker's content
        if current_text and current_speaker:
            text = " ".join(current_text)
            story.append(Paragraph(
                f"<b>{current_speaker}:</b> {text}",
                speaker_style
            ))
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph("_" * 70, styles['Normal']))
        story.append(Spacer(1, 10))
        story.append(Paragraph(
            "<i>Gegenereerd door Transcriptie PoC - Raad van State</i>",
            meta_style
        ))
        
        # Build PDF
        doc.build(story)
        buffer.seek(0)
        
        return buffer
